#!/usr/bin/env python3
"""
Create CompLaB3D Tutorial Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_tutorial():
    doc = Document()

    # Title
    title = doc.add_heading('CompLaB3D Technical Tutorial', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Comprehensive Guide to Biotic/Abiotic Kinetics, Cellular Automata, and Equilibrium Solver')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction',
        '2. Biotic Kinetics (Monod Model)',
        '3. Abiotic Kinetics',
        '4. Cellular Automata (CA) Biofilm Expansion',
        '5. Equilibrium Solver (Anderson Acceleration)',
        '6. XML Configuration Guide',
        '7. Test Cases and Validation'
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ===========================================
    # SECTION 1: INTRODUCTION
    # ===========================================
    doc.add_heading('1. Introduction', level=1)

    doc.add_paragraph(
        'CompLaB3D is a 3D reactive transport simulation framework using the Lattice Boltzmann Method (LBM). '
        'It supports both biotic (with microorganisms) and abiotic (without microorganisms) simulations.'
    )

    doc.add_heading('1.1 Simulation Modes', level=2)

    # Biotic mode
    p = doc.add_paragraph()
    p.add_run('Biotic Mode (biotic_mode = true):').bold = True
    doc.add_paragraph('- Simulates microorganism growth and activity', style='List Bullet')
    doc.add_paragraph('- Uses Monod kinetics for substrate consumption', style='List Bullet')
    doc.add_paragraph('- Biofilm expansion via Cellular Automata', style='List Bullet')
    doc.add_paragraph('- Biomass grows, decays, and spreads', style='List Bullet')

    # Abiotic mode
    p = doc.add_paragraph()
    p.add_run('Abiotic Mode (biotic_mode = false):').bold = True
    doc.add_paragraph('- No microorganisms present', style='List Bullet')
    doc.add_paragraph('- Pure transport (diffusion/advection) OR', style='List Bullet')
    doc.add_paragraph('- Transport + chemical reactions (abiotic kinetics)', style='List Bullet')
    doc.add_paragraph('- Useful for geochemical modeling', style='List Bullet')

    doc.add_page_break()

    # ===========================================
    # SECTION 2: BIOTIC KINETICS
    # ===========================================
    doc.add_heading('2. Biotic Kinetics (Monod Model)', level=1)

    doc.add_paragraph(
        'When microorganisms are present (biotic_mode = true), substrate consumption and biomass growth '
        'follow Monod kinetics. This is the standard model for microbial growth.'
    )

    doc.add_heading('2.1 The Monod Equation', level=2)

    doc.add_paragraph('The specific growth rate is calculated as:')

    # Monod equation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('mu = mu_max * C / (Ks + C)')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph('Where:')
    doc.add_paragraph('- mu = specific growth rate (1/s)', style='List Bullet')
    doc.add_paragraph('- mu_max = maximum specific growth rate (1/s)', style='List Bullet')
    doc.add_paragraph('- C = substrate concentration (kg/m3)', style='List Bullet')
    doc.add_paragraph('- Ks = half-saturation constant (kg/m3)', style='List Bullet')

    doc.add_heading('2.2 Biomass Change', level=2)

    doc.add_paragraph('The biomass changes according to:')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('dB/dt = (mu - k_decay) * B')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph('Where:')
    doc.add_paragraph('- B = biomass concentration (kg/m3)', style='List Bullet')
    doc.add_paragraph('- k_decay = decay rate constant (1/s)', style='List Bullet')

    doc.add_heading('2.3 Substrate Consumption', level=2)

    doc.add_paragraph('Substrate is consumed proportionally to biomass growth:')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('dC/dt = -(1/Y) * mu * B')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph('Where:')
    doc.add_paragraph('- Y = yield coefficient (kg biomass / kg substrate)', style='List Bullet')

    doc.add_heading('2.4 Code Implementation', level=2)

    doc.add_paragraph('In defineKinetics.hh, the Monod kinetics are implemented:')

    # Code block
    code = '''// Monod kinetics calculation
T mu = mumax[iM] * C[iS] / (Ks[iM] + C[iS]);

// Biomass change
dB = (mu - k_decay[iM]) * B[iM] * dt;

// Substrate consumption
dC = -(1.0/Y[iM]) * mu * B[iM] * dt;'''

    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('2.5 XML Configuration for Biotic Mode', level=2)

    xml_code = '''<microbe>
    <substrate>substrate_A</substrate>
    <mumax>1.0e-5</mumax>      <!-- Max growth rate (1/s) -->
    <Ks>0.001</Ks>              <!-- Half-saturation (kg/m3) -->
    <Y>0.5</Y>                  <!-- Yield coefficient -->
    <kdecay>1.0e-6</kdecay>     <!-- Decay rate (1/s) -->
    <b0_biofilm>0.1</b0_biofilm><!-- Initial biomass (kg/m3) -->
</microbe>'''

    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(xml_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_page_break()

    # ===========================================
    # SECTION 3: ABIOTIC KINETICS
    # ===========================================
    doc.add_heading('3. Abiotic Kinetics', level=1)

    doc.add_paragraph(
        'When biotic_mode = false and enable_abiotic_kinetics = true, chemical reactions occur '
        'without microorganisms. This is useful for modeling purely chemical systems.'
    )

    doc.add_heading('3.1 Reaction Types', level=2)

    # Type 1: First-order decay
    doc.add_heading('3.1.1 First-Order Decay', level=3)

    p = doc.add_paragraph()
    p.add_run('Reaction: A -> products').bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('dC_A/dt = -k * C_A')
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph('Analytical solution: C_A(t) = C_A(0) * exp(-k*t)')

    code = '''// In defineAbioticKinetics.hh:
T k = 0.001;  // decay rate (1/s)
dC[0] = -k * C[0] * dt;  // substrate A'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Type 2: Bimolecular
    doc.add_heading('3.1.2 Bimolecular Reaction', level=3)

    p = doc.add_paragraph()
    p.add_run('Reaction: A + B -> C').bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('rate = k * C_A * C_B')
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph('dC_A/dt = dC_B/dt = -rate')
    doc.add_paragraph('dC_C/dt = +rate')

    code = '''// In defineAbioticKinetics.hh:
T k = 100.0;  // rate constant (m3/kg/s)
T rate = k * C[0] * C[1] * dt;
dC[0] = -rate;  // A consumed
dC[1] = -rate;  // B consumed
dC[2] = +rate;  // C produced'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Type 3: Reversible
    doc.add_heading('3.1.3 Reversible Reaction', level=3)

    p = doc.add_paragraph()
    p.add_run('Reaction: A <-> B').bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('net_rate = k_forward * C_A - k_reverse * C_B')
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph('At equilibrium: K_eq = k_forward / k_reverse = C_B / C_A')

    code = '''// In defineAbioticKinetics.hh:
T k_f = 0.002;   // forward rate
T k_r = 0.001;   // reverse rate (K_eq = 2)
T net_rate = (k_f * C[0] - k_r * C[1]) * dt;
dC[0] = -net_rate;  // A
dC[1] = +net_rate;  // B'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Type 4: Sequential
    doc.add_heading('3.1.4 Sequential Decay Chain', level=3)

    p = doc.add_paragraph()
    p.add_run('Reaction: A -> B -> C').bold = True

    code = '''// In defineAbioticKinetics.hh:
T k1 = 0.002;  // A -> B rate
T k2 = 0.001;  // B -> C rate
dC[0] = -k1 * C[0] * dt;           // A decays
dC[1] = (k1*C[0] - k2*C[1]) * dt;  // B: formed and decays
dC[2] = k2 * C[1] * dt;            // C accumulates'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('3.2 XML Configuration for Abiotic Mode', level=2)

    xml_code = '''<options>
    <biotic_mode>false</biotic_mode>
    <enable_kinetics>true</enable_kinetics>
    <enable_abiotic_kinetics>true</enable_abiotic_kinetics>
</options>

<!-- No <microbe> sections needed -->
<!-- Define reactions in defineAbioticKinetics.hh -->'''

    p = doc.add_paragraph()
    run = p.add_run(xml_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_page_break()

    # ===========================================
    # SECTION 4: CELLULAR AUTOMATA
    # ===========================================
    doc.add_heading('4. Cellular Automata (CA) Biofilm Expansion', level=1)

    doc.add_paragraph(
        'When biomass exceeds a threshold, it spreads to neighboring cells using Cellular Automata rules. '
        'CompLaB3D implements two methods: "fraction" and "half".'
    )

    doc.add_heading('4.1 When Does Expansion Occur?', level=2)

    doc.add_paragraph('Expansion is triggered when:')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('B_cell > B_max (maximum biomass capacity)')
    run.bold = True

    doc.add_heading('4.2 The "fraction" Method', level=2)

    doc.add_paragraph('Excess biomass is distributed equally to all valid neighbors:')

    doc.add_paragraph('1. Calculate excess: excess = B_cell - B_max', style='List Number')
    doc.add_paragraph('2. Find valid neighbors (pore or existing biofilm)', style='List Number')
    doc.add_paragraph('3. Distribute: each neighbor gets excess / N_neighbors', style='List Number')
    doc.add_paragraph('4. Parent cell keeps B_max', style='List Number')

    code = '''// In complab3d_processors_part2.hh:
T excess = B_cell - B_max;
int N = valid_neighbors.size();
T share = excess / N;
for (each neighbor) {
    B_neighbor += share;
}
B_cell = B_max;'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('4.3 The "half" Method', level=2)

    doc.add_paragraph('Parent cell splits, giving half to ONE random neighbor:')

    doc.add_paragraph('1. Calculate excess: excess = B_cell - B_max', style='List Number')
    doc.add_paragraph('2. Find valid neighbors', style='List Number')
    doc.add_paragraph('3. Select ONE random neighbor', style='List Number')
    doc.add_paragraph('4. Transfer: neighbor gets B_cell / 2', style='List Number')
    doc.add_paragraph('5. Parent keeps B_cell / 2', style='List Number')

    code = '''// In complab3d_processors_part2.hh:
T excess = B_cell - B_max;
int random_idx = rand() % valid_neighbors.size();
neighbor = valid_neighbors[random_idx];
B_neighbor += B_cell / 2;
B_cell = B_cell / 2;'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('4.4 XML Configuration', level=2)

    xml_code = '''<microbe>
    <bmax>100.0</bmax>           <!-- Threshold for expansion -->
    <ca_expansion>fraction</ca_expansion>  <!-- or "half" -->
</microbe>'''

    p = doc.add_paragraph()
    run = p.add_run(xml_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('4.5 Comparison of Methods', level=2)

    # Create comparison table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Header
    hdr = table.rows[0].cells
    hdr[0].text = 'Aspect'
    hdr[1].text = 'fraction'
    hdr[2].text = 'half'

    # Data
    data = [
        ['Distribution', 'Equal to all neighbors', 'All to one neighbor'],
        ['Growth pattern', 'Smooth, uniform', 'Irregular, branching'],
        ['Use case', 'Dense biofilms', 'Finger-like structures']
    ]

    for i, row_data in enumerate(data):
        row = table.rows[i+1].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data

    doc.add_page_break()

    # ===========================================
    # SECTION 5: EQUILIBRIUM SOLVER
    # ===========================================
    doc.add_heading('5. Equilibrium Solver (Anderson Acceleration)', level=1)

    doc.add_paragraph(
        'CompLaB3D uses the Picard iteration with Anderson Acceleration to solve for steady-state '
        'concentration fields. This is faster than simple iteration.'
    )

    doc.add_heading('5.1 The Problem', level=2)

    doc.add_paragraph('We want to find the steady-state solution where:')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('dC/dt = 0  (concentrations stop changing)')
    run.bold = True

    doc.add_heading('5.2 Picard Iteration (Fixed-Point)', level=2)

    doc.add_paragraph('Basic approach - iterate until convergence:')

    code = '''C_new = G(C_old)  // Apply one LBM step
repeat until ||C_new - C_old|| < tolerance'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_paragraph('Problem: Simple Picard iteration can be SLOW (thousands of iterations).')

    doc.add_heading('5.3 Anderson Acceleration', level=2)

    doc.add_paragraph(
        'Anderson Acceleration speeds up convergence by using information from previous iterations '
        'to make better guesses. It combines the last m iterations optimally.'
    )

    doc.add_paragraph('Key idea:')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('C_new = sum(alpha_i * G(C_i))  where sum(alpha_i) = 1')
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph('The weights alpha_i are chosen to minimize the residual.')

    doc.add_heading('5.4 Algorithm Steps', level=2)

    doc.add_paragraph('1. Store history of last m iterations', style='List Number')
    doc.add_paragraph('2. Build matrix F of residuals: F_k = G(C_k) - C_k', style='List Number')
    doc.add_paragraph('3. Solve least-squares problem for optimal weights', style='List Number')
    doc.add_paragraph('4. Compute accelerated update', style='List Number')
    doc.add_paragraph('5. Apply relaxation: C = beta * C_accelerated + (1-beta) * C_old', style='List Number')

    doc.add_heading('5.5 Code Implementation', level=2)

    code = '''// In complab3d_equilibrium.hh:

// Store iteration history
std::vector<MultiScalarField3D<T>*> X_history;  // solutions
std::vector<MultiScalarField3D<T>*> F_history;  // residuals

// Anderson mixing
for (int k = 0; k < max_iter; k++) {
    // 1. Compute G(x_k) - one LBM step
    G_k = applyLBMStep(X_k);

    // 2. Compute residual
    F_k = G_k - X_k;

    // 3. Check convergence
    if (norm(F_k) < tolerance) break;

    // 4. Anderson acceleration (if enough history)
    if (k >= m) {
        // Build matrix and solve least-squares
        alpha = solveLeastSquares(F_history);

        // Accelerated update
        X_new = sum(alpha[i] * G_history[i]);
    } else {
        X_new = G_k;  // Simple Picard
    }

    // 5. Relaxation
    X_k = beta * X_new + (1-beta) * X_k;
}'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('5.6 XML Configuration', level=2)

    xml_code = '''<options>
    <solver_type>equilibrium</solver_type>
</options>

<equilibrium>
    <tolerance>1e-6</tolerance>     <!-- Convergence criterion -->
    <max_iterations>10000</max_iterations>
    <anderson_depth>5</anderson_depth>  <!-- History size m -->
    <relaxation>0.5</relaxation>    <!-- Mixing parameter beta -->
</equilibrium>'''

    p = doc.add_paragraph()
    run = p.add_run(xml_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('5.7 Convergence Comparison', level=2)

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Typical Iterations'
    hdr[2].text = 'Convergence'

    data = [
        ['Simple Picard', '5000-50000', 'Linear (slow)'],
        ['Anderson (m=5)', '500-2000', 'Superlinear (fast)']
    ]

    for i, row_data in enumerate(data):
        row = table.rows[i+1].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data

    doc.add_page_break()

    # ===========================================
    # SECTION 6: XML CONFIGURATION
    # ===========================================
    doc.add_heading('6. XML Configuration Guide', level=1)

    doc.add_heading('6.1 Complete Biotic Example', level=2)

    xml_code = '''<?xml version="1.0" ?>
<complab>
    <options>
        <biotic_mode>true</biotic_mode>
        <enable_kinetics>true</enable_kinetics>
    </options>

    <geometry>
        <file>geometry.dat</file>
        <nx>100</nx>
        <ny>50</ny>
        <nz>50</nz>
    </geometry>

    <simulation>
        <dt>0.1</dt>
        <total_time>86400</total_time>  <!-- 1 day -->
        <output_interval>3600</output_interval>
    </simulation>

    <substrate name="oxygen">
        <diffusion>
            <in_water>2.1e-9</in_water>
            <in_biofilm>1.0e-9</in_biofilm>
        </diffusion>
        <initial_concentration>0.008</initial_concentration>
        <boundary>
            <left type="dirichlet">0.008</left>
            <right type="neumann">0.0</right>
        </boundary>
    </substrate>

    <microbe name="aerobic_bacteria">
        <substrate>oxygen</substrate>
        <mumax>1.0e-5</mumax>
        <Ks>0.001</Ks>
        <Y>0.5</Y>
        <kdecay>1.0e-6</kdecay>
        <b0_biofilm>0.1</b0_biofilm>
        <bmax>100.0</bmax>
        <ca_expansion>fraction</ca_expansion>
    </microbe>
</complab>'''

    p = doc.add_paragraph()
    run = p.add_run(xml_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

    doc.add_heading('6.2 Complete Abiotic Example', level=2)

    xml_code = '''<?xml version="1.0" ?>
<complab>
    <options>
        <biotic_mode>false</biotic_mode>
        <enable_kinetics>true</enable_kinetics>
        <enable_abiotic_kinetics>true</enable_abiotic_kinetics>
    </options>

    <geometry>
        <file>geometry.dat</file>
        <nx>100</nx>
        <ny>30</ny>
        <nz>30</nz>
    </geometry>

    <simulation>
        <dt>0.1</dt>
        <total_time>10000</total_time>
        <output_interval>1000</output_interval>
    </simulation>

    <!-- Reactant A -->
    <substrate name="species_A">
        <diffusion>
            <in_water>1.0e-9</in_water>
            <in_biofilm>1.0e-9</in_biofilm>
        </diffusion>
        <initial_concentration>1.0</initial_concentration>
        <boundary>
            <left type="dirichlet">1.0</left>
            <right type="neumann">0.0</right>
        </boundary>
    </substrate>

    <!-- Product B (from A decay) -->
    <substrate name="species_B">
        <diffusion>
            <in_water>1.0e-9</in_water>
            <in_biofilm>1.0e-9</in_biofilm>
        </diffusion>
        <initial_concentration>0.0</initial_concentration>
        <boundary>
            <left type="dirichlet">0.0</left>
            <right type="neumann">0.0</right>
        </boundary>
    </substrate>

    <!-- NO <microbe> sections in abiotic mode -->
    <!-- Reactions defined in defineAbioticKinetics.hh -->
</complab>'''

    p = doc.add_paragraph()
    run = p.add_run(xml_code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

    doc.add_page_break()

    # ===========================================
    # SECTION 7: TEST CASES
    # ===========================================
    doc.add_heading('7. Test Cases and Validation', level=1)

    doc.add_heading('7.1 Abiotic Test Cases', level=2)

    # Test 1
    doc.add_heading('Test 1: Pure Diffusion', level=3)
    doc.add_paragraph('Purpose: Verify diffusion without reactions')
    doc.add_paragraph('Setup: Species A diffuses from left (C=1.0) to right')
    doc.add_paragraph('Expected: Linear concentration profile at steady state')

    # Test 2
    doc.add_heading('Test 2: First-Order Decay (A -> products)', level=3)
    doc.add_paragraph('Purpose: Verify first-order reaction kinetics')
    doc.add_paragraph('Rate law: dC/dt = -k*C with k = 0.001 1/s')
    doc.add_paragraph('Validation: C(t) = C0 * exp(-k*t)')
    doc.add_paragraph('At t=1000s: C should be ~0.368 * C0')

    # Test 3
    doc.add_heading('Test 3: Bimolecular (A + B -> C)', level=3)
    doc.add_paragraph('Purpose: Verify second-order reaction kinetics')
    doc.add_paragraph('Rate law: rate = k * C_A * C_B')
    doc.add_paragraph('Validation: Mass balance A + B + C = constant')

    # Test 4
    doc.add_heading('Test 4: Reversible (A <-> B)', level=3)
    doc.add_paragraph('Purpose: Verify equilibrium is reached')
    doc.add_paragraph('Rate law: net = k_f*A - k_r*B')
    doc.add_paragraph('Validation: At equilibrium, B/A = k_f/k_r = K_eq')

    # Test 5
    doc.add_heading('Test 5: Decay Chain (A -> B -> C)', level=3)
    doc.add_paragraph('Purpose: Verify sequential reactions')
    doc.add_paragraph('Validation: B shows characteristic rise-then-fall; C accumulates')

    doc.add_heading('7.2 Running Test Cases', level=2)

    code = '''# Navigate to test directory
cd test_cases/abiotic

# Copy test files
cp test2_first_order_decay.xml CompLaB.xml
cp defineAbioticKinetics_test2.hh ../../src/defineAbioticKinetics.hh

# Recompile and run
cd ../../src
make clean && make
cd ..
./complab

# Check results in output/'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    doc.add_heading('7.3 Validation Criteria', level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    hdr[0].text = 'Test'
    hdr[1].text = 'Check'
    hdr[2].text = 'Tolerance'

    data = [
        ['Pure Diffusion', 'Linear profile', '< 1%'],
        ['First-Order', 'C(t)/C0 = exp(-kt)', '< 5%'],
        ['Bimolecular', 'Mass conserved', '< 1%'],
        ['Reversible', 'B/A = K_eq', '< 5%'],
        ['Decay Chain', 'A+B+C = constant', '< 1%']
    ]

    for i, row_data in enumerate(data):
        row = table.rows[i+1].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data

    # Save document
    doc.save('/home/user/Complab3D-Biotic-Kinetic-AndersopnNR/docs/CompLaB3D_Tutorial.docx')
    print("Tutorial document created: docs/CompLaB3D_Tutorial.docx")

if __name__ == '__main__':
    create_tutorial()
