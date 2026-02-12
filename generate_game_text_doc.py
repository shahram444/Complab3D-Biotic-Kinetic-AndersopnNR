"""Generate a Word document with all game dialogues and text from ARKE: Guardians of Earth."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ─── TITLE PAGE ─────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ARKE: Guardians of Earth')
run.bold = True
run.font.size = Pt(26)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Complete Game Dialogues, Texts, and Information')
run.font.size = Pt(16)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Based on CompLaB3D Research\nPore-Scale Biogeochemical Reactive Transport\nUniversity of Georgia')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ─── TABLE OF CONTENTS (manual) ────────────────────────────────────────────
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Boot Screen',
    '2. Title Screen',
    '3. Opening Cutscene (Intro Narrative)',
    '4. Chapter 1: The Soil Frontier',
    '   4.1  Chapter 1 Introduction',
    '   4.2  Level 1 – First Breath (Mission Briefing)',
    '   4.3  Level 1 – Level Complete',
    '   4.4  Level 1 – Science Discovery: The Redox Ladder',
    '   4.5  Level 2 – Roots of Life (Mission Briefing)',
    '   4.6  Level 2 – Level Complete',
    '   4.7  Level 2 – Science Discovery: Methanotrophy',
    '5. Chapter 2: The Deep Sediment',
    '   5.1  Chapter 2 Introduction',
    '   5.2  Level 3 – Into the Depths (Mission Briefing)',
    '   5.3  Level 3 – Level Complete',
    '   5.4  Level 3 – Science Discovery: Denitrification',
    '   5.5  Level 4 – The Hungry Dark (Mission Briefing)',
    '   5.6  Level 4 – Level Complete',
    '   5.7  Level 4 – Science Discovery: Diffusion vs Advection',
    '6. Chapter 3: The Methane Seeps',
    '   6.1  Chapter 3 Introduction',
    '   6.2  Level 5 – The Methane Vents (Mission Briefing)',
    '   6.3  Level 5 – Level Complete',
    '   6.4  Level 5 – Science Discovery: Monod Kinetics',
    '   6.5  Level 6 – Vent Guardians (Mission Briefing)',
    '   6.6  Level 6 – Level Complete',
    '   6.7  Level 6 – Science Discovery: Biofilm Formation (CA)',
    '7. Chapter 4: The Permafrost Edge',
    '   7.1  Chapter 4 Introduction',
    '   7.2  Level 7 – Thawing Grounds (Mission Briefing)',
    '   7.3  Level 7 – Level Complete',
    '   7.4  Level 7 – Science Discovery: Anaerobic Methane Oxidation',
    '   7.5  Level 8 – The Great Thaw (Mission Briefing)',
    '   7.6  Level 8 – Level Complete',
    '   7.7  Level 8 – Science Discovery: Permafrost Carbon Feedback',
    '8. Chapter 5: The Hydrothermal Realm',
    '   8.1  Chapter 5 Introduction',
    '   8.2  Level 9 – The Abyss (Mission Briefing)',
    '   8.3  Level 9 – Level Complete',
    '   8.4  Level 9 – Science Discovery: Hydrothermal Vent Chemistry',
    '   8.5  Level 10 – Earth\'s Last Stand (Mission Briefing)',
    '   8.6  Level 10 – Level Complete',
    '   8.7  Level 10 – Science Discovery: You Are Earth\'s Climate Shield',
    '9. Victory Screen',
    '10. Game Over Screen',
    '11. Tutorial Hints',
    '12. Pause Menu',
    '13. HUD Elements and In-Game Labels',
    '14. Substrate Information (Redox Ladder)',
    '15. Science Parameters',
    '16. Environment Descriptions',
    '17. Level Definitions',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1: BOOT SCREEN
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Boot Screen', level=1)
doc.add_paragraph('The boot screen appears for 2 seconds when the game starts.')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('CompLaB3D').bold = True
doc.add_paragraph('presents')
doc.add_paragraph()
doc.add_paragraph('A Game Based on Real Science')

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2: TITLE SCREEN
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Title Screen', level=1)
doc.add_paragraph('The title screen shows the game name, animated character sprite, and prompts.')
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('ARKE')
run.bold = True
run.font.size = Pt(18)
doc.add_paragraph('Guardians of Earth')
doc.add_paragraph()
doc.add_paragraph('PRESS ENTER')
doc.add_paragraph()
doc.add_paragraph('Based on CompLaB3D Research')
doc.add_paragraph('Pore-Scale Reactive Transport')
doc.add_paragraph('University of Georgia')

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3: OPENING CUTSCENE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Opening Cutscene (Intro Narrative)', level=1)
doc.add_paragraph(
    'The opening cutscene plays as a visual novel-style dialogue sequence after the player '
    'presses ENTER on the title screen. Two characters speak: ARCHAEON PRIME (the Elder mentor) '
    'and ARKE (the player character). Narration lines have no speaker. '
    'Text appears with a typewriter effect.'
)
doc.add_paragraph()

cutscene = [
    {"speaker": "", "text": "Deep within the pore space...\nbetween grains of ancient rock..."},
    {"speaker": "ELDER", "text": "Young one... wake up."},
    {"speaker": "ARKE", "text": "W-where am I? What is this place?"},
    {"speaker": "ELDER", "text": "You are in the subsurface. Between\ngrains of soil, far from light.\nThis is our world."},
    {"speaker": "ELDER", "text": "I am ARCHAEON PRIME.\nElder of the methanotrophic archaea.\nI have guarded these pores for eons."},
    {"speaker": "ARKE", "text": "Why is everything so dark?\nWhat's happening?"},
    {"speaker": "ELDER", "text": "Methane. CH4. A greenhouse gas\n80 times more powerful than CO2.\nIt rises from the deep below."},
    {"speaker": "ELDER", "text": "And N2O... nitrous oxide.\n300 times more dangerous than CO2.\nThey seep from thawing soil."},
    {"speaker": "ELDER", "text": "For billions of years, we microbes\nhave consumed these gases.\nWe are Earth's invisible shield."},
    {"speaker": "ARKE", "text": "What must I do?"},
    {"speaker": "ELDER", "text": "Eat substrates. Grow your biomass.\nWhen your growth bar fills -\npress SPACE to DIVIDE!"},
    {"speaker": "ELDER", "text": "Place colonies across the pore network.\nBuild a biofilm to filter the gas.\nThat is the Cellular Automata way."},
    {"speaker": "ELDER", "text": "But beware... rival microbes roam\nthese pores. They will steal\nyour nutrients."},
    {"speaker": "ELDER", "text": "And the pore geometry itself is\nyour enemy. Dead-end pores\nmean starvation and death."},
    {"speaker": "ARKE", "text": "I won't let Earth down."},
    {"speaker": "ELDER", "text": "Remember the redox ladder:\nO2 gives most energy, CH4 the least.\nBut eating CH4 saves the planet!"},
    {"speaker": "ELDER", "text": "Now go, young one.\nThe Soil Frontier awaits.\nEarth's fate is in your pseudopods."},
]

for i, line in enumerate(cutscene):
    speaker = line["speaker"]
    text = line["text"].replace("\n", "\n")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    if speaker == "":
        run = p.add_run(f'[Narration]')
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 140)
    elif speaker == "ELDER":
        run = p.add_run('ARCHAEON PRIME:')
        run.bold = True
        run.font.color.rgb = RGBColor(70, 70, 210)
    elif speaker == "ARKE":
        run = p.add_run('ARKE:')
        run.bold = True
        run.font.color.rgb = RGBColor(40, 207, 176)

    # Add the text on new lines
    for tline in text.split("\n"):
        p2 = doc.add_paragraph(tline)
        p2.paragraph_format.left_indent = Inches(0.5)
        p2.paragraph_format.space_after = Pt(1)
        p2.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# Helper function for mission briefings
# ═══════════════════════════════════════════════════════════════════════════
def add_mission_briefing(doc, briefing_lines, level_num, level_title):
    """Add a mission briefing section."""
    doc.add_heading(f'Level {level_num} – {level_title} (Mission Briefing)', level=3)
    doc.add_paragraph(
        f'Shown before Level {level_num} begins. The Elder briefs ARKE on the upcoming mission.'
    )
    doc.add_paragraph()

    for line in briefing_lines:
        speaker = line["speaker"]
        text = line["text"]

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

        if speaker == "ELDER":
            run = p.add_run('ARCHAEON PRIME:')
            run.bold = True
            run.font.color.rgb = RGBColor(70, 70, 210)
        elif speaker == "ARKE":
            run = p.add_run('ARKE:')
            run.bold = True
            run.font.color.rgb = RGBColor(40, 207, 176)

        for tline in text.split("\n"):
            p2 = doc.add_paragraph(tline)
            p2.paragraph_format.left_indent = Inches(0.5)
            p2.paragraph_format.space_after = Pt(1)
            p2.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()


def add_level_complete(doc, level_num, level_title):
    """Add a level complete screen section."""
    doc.add_heading(f'Level {level_num} – Level Complete', level=3)
    doc.add_paragraph('Displayed when the player places enough colonies to meet the level goal.')
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('LEVEL COMPLETE!')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(94, 207, 94)
    doc.add_paragraph()
    doc.add_paragraph('Substrates eaten: [count]')
    doc.add_paragraph('Methane consumed: [count]')
    doc.add_paragraph('N2O prevented: [count]')
    doc.add_paragraph('Colonies placed: [count]')
    doc.add_paragraph('Score: [points]')
    doc.add_paragraph()
    doc.add_paragraph('PRESS ENTER')
    doc.add_paragraph()


def add_science_fact(doc, level_num, fact):
    """Add a science discovery popup section."""
    doc.add_heading(f'Level {level_num} – Science Discovery: {fact["title"]}', level=3)
    doc.add_paragraph('Shown after the Level Complete screen. Teaches real science.')
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('SCIENCE DISCOVERY')
    run.bold = True
    run.font.color.rgb = RGBColor(79, 163, 240)

    p = doc.add_paragraph()
    run = p.add_run(fact["title"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(200, 170, 50)

    doc.add_paragraph()
    for tline in fact["text"].split("\n"):
        p = doc.add_paragraph(tline)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

    doc.add_paragraph()
    doc.add_paragraph('PRESS ENTER')
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# Mission briefings data
# ═══════════════════════════════════════════════════════════════════════════
mission_briefs = [
    # Level 1: First Breath
    [
        {"speaker": "ELDER", "text": "Your first mission, young one.\nThis is the Soil Frontier -\nshallow soil beneath a meadow."},
        {"speaker": "ELDER", "text": "Oxygen flows from above.\nEat substrates to survive.\nPlace 3 colonies to secure this zone."},
        {"speaker": "ARKE", "text": "I'm ready, Archaeon Prime!"},
    ],
    # Level 2: Roots of Life
    [
        {"speaker": "ELDER", "text": "You've grown stronger.\nBut now a rival microbe lurks\nin these pores. Be vigilant."},
        {"speaker": "ELDER", "text": "It will steal your food.\nMove fast, eat faster.\nPlace 5 colonies this time."},
    ],
    # Level 3: Into the Depths
    [
        {"speaker": "ELDER", "text": "We descend into the Deep Sediment.\nOxygen cannot reach here.\nOnly iron and nitrate remain."},
        {"speaker": "ELDER", "text": "The passages are narrow and dark.\nTwo rivals hunt these corridors.\nConserve your energy carefully."},
        {"speaker": "ARKE", "text": "No oxygen? How do I survive?"},
        {"speaker": "ELDER", "text": "Use the redox ladder, child.\nNitrate and iron will sustain you.\nAnd never stop eating methane!"},
    ],
    # Level 4: The Hungry Dark
    [
        {"speaker": "ELDER", "text": "Deeper still. The sediment\ntightens around us.\nThree rivals compete for scraps."},
        {"speaker": "ELDER", "text": "Manganese joins your options now.\nBut food is scarce - every\nmolecule matters here."},
    ],
    # Level 5: The Methane Vents
    [
        {"speaker": "ELDER", "text": "Now we enter the Methane Seeps!\nCH4 erupts from volcanic vents below.\nThis is where our kind truly shines."},
        {"speaker": "ELDER", "text": "Beware the toxic zones - they\nburn through your cell membrane.\nAvoid the purple areas!"},
        {"speaker": "ARKE", "text": "The methane... I can feel it rising."},
        {"speaker": "ELDER", "text": "Consume it! Every molecule of CH4\nyou eat prevents global warming.\nYou are Earth's climate shield!"},
    ],
    # Level 6: Vent Guardians
    [
        {"speaker": "ELDER", "text": "The vents grow stronger.\nMore methane, more toxins,\nmore rivals. Stay focused."},
        {"speaker": "ELDER", "text": "Place 8 colonies to build\na biofilm wall against\nthe rising greenhouse gases."},
    ],
    # Level 7: Thawing Grounds
    [
        {"speaker": "ELDER", "text": "The Permafrost Edge... the ice\nis melting. Ancient carbon\nunlocks after millennia."},
        {"speaker": "ELDER", "text": "The flow here is fast and chaotic.\nUse SHIFT to ride the currents.\nIt will save your energy."},
        {"speaker": "ARKE", "text": "The water moves so fast here!"},
        {"speaker": "ELDER", "text": "The melting permafrost releases\n1,500 gigatons of trapped carbon.\nWe are the last line of defense."},
    ],
    # Level 8: The Great Thaw
    [
        {"speaker": "ELDER", "text": "This is critical. The thaw\naccelerates. Four rivals compete.\nMethane pulses are massive."},
        {"speaker": "ELDER", "text": "Build 8 colonies. Create a\nbiofilm barrier. The planet's\nfuture depends on us."},
    ],
    # Level 9: The Abyss
    [
        {"speaker": "ELDER", "text": "The Hydrothermal Realm.\nDeep ocean vents blast superheated\nfluid through mineral chimneys."},
        {"speaker": "ELDER", "text": "The full redox ladder is here -\nsulfate, iron, manganese, methane.\nUse everything available to you."},
        {"speaker": "ARKE", "text": "It's so hot... and hostile."},
        {"speaker": "ELDER", "text": "Only the strongest archaea\nsurvive here. But you have\ngrown beyond my expectations."},
    ],
    # Level 10: Earth's Last Stand
    [
        {"speaker": "ELDER", "text": "This is it. The final stand.\nEvery greenhouse gas that escapes\nthese vents warms the planet."},
        {"speaker": "ELDER", "text": "Five rivals. Toxic zones everywhere.\n12 colonies needed to seal\nthe vent network permanently."},
        {"speaker": "ARKE", "text": "I won't let you down, Elder."},
        {"speaker": "ELDER", "text": "You have become a true Guardian\nof Earth. Now show the world\nwhat one archaea can do!"},
    ],
]

# Science facts data
science_facts = [
    {"title": "The Redox Ladder",
     "text": "Microbes harvest energy by transferring electrons\nfrom donors to acceptors. The 'redox ladder' ranks\nthese reactions by energy yield:\n\n  O2       -> 818 kJ/mol  (most energy)\n  NO3-     -> 649 kJ/mol\n  Mn(IV)   -> 558 kJ/mol\n  Fe(III)  -> 334 kJ/mol\n  SO42-    -> 152 kJ/mol\n  CO2/CH4  ->  31 kJ/mol  (least energy)\n\nMicrobes prefer the highest available energy source."},
    {"title": "Methanotrophy",
     "text": "Methanotrophs consume methane (CH4) before\nit reaches the atmosphere. Without them,\natmospheric CH4 would be 10-100x higher.\n\n  CH4 + 2O2 -> CO2 + 2H2O\n\nThey are Earth's methane filter,\noperating in soils, wetlands, and ocean\nsediments worldwide."},
    {"title": "Denitrification",
     "text": "Denitrifying microbes convert nitrate (NO3-)\nto harmless nitrogen gas (N2).\n\n  2NO3- + 10e- + 12H+ -> N2 + 6H2O\n\nWithout them, nitrous oxide (N2O) - a\ngreenhouse gas 300x stronger than CO2 -\nwould accumulate in the atmosphere."},
    {"title": "Diffusion vs Advection",
     "text": "Nutrients reach microbes two ways:\n\n  Advection: carried by flowing water\n    (fast, directional)\n  Diffusion: random molecular motion\n    (slow, spreads in all directions)\n\nThe Peclet number (Pe = uL/D) tells us\nwhich dominates. In deep sediments,\nPe << 1: diffusion rules."},
    {"title": "Monod Kinetics",
     "text": "Microbial growth follows the Monod equation:\n\n  mu = mu_max * C / (Ks + C)\n\nWhen substrate C >> Ks: growth is maximal\nWhen C << Ks: growth nearly stops\nAt C = Ks: growth is half-maximum\n\nThis is why position matters - cells\nnear flow paths get more food."},
    {"title": "Biofilm Formation (CA)",
     "text": "When biomass exceeds B_max (100 kg/m3),\nexcess is pushed to neighboring pores.\nThis Cellular Automata model creates\nrealistic biofilm growth patterns.\n\nIn CompLaB3D, a distance field guides\ngrowth toward open pore space,\nmimicking real biofilm expansion."},
    {"title": "Anaerobic Methane Oxidation",
     "text": "In oxygen-free zones, archaea partner\nwith sulfate-reducing bacteria:\n\n  CH4 + SO42- -> HCO3- + HS- + H2O\n\nThis 'anaerobic oxidation of methane'\n(AOM) consumes ~90% of oceanic methane\nbefore it reaches the water column."},
    {"title": "Permafrost Carbon Feedback",
     "text": "Arctic permafrost stores ~1,500 Gt carbon.\nAs it thaws, microbes decompose this\norganic matter, releasing CH4 and CO2.\n\nBut methanotrophs in the thaw layer\nact as a biological filter, consuming\nmuch of this methane before release.\n\nYour colonies ARE that filter."},
    {"title": "Hydrothermal Vent Chemistry",
     "text": "Deep-sea vents release fluids at 300-400C\nrich in H2S, Fe2+, Mn2+, and CH4.\n\nChemosynthetic microbes harvest energy\nfrom these chemicals, forming the base\nof entire ecosystems without sunlight.\n\nThe full redox ladder operates here\nin centimeters of sediment."},
    {"title": "You Are Earth's Climate Shield",
     "text": "Subsurface microbes process ~600 Tg CH4/yr,\npreventing most geological methane from\nreaching the atmosphere.\n\nWithout microbial methane consumption,\nEarth's temperature would be\nsignificantly higher.\n\nEvery colony you built represents\na real climate defense mechanism.\n\nCongratulations, Guardian of Earth."},
]

# World intros data
world_intros = [
    {"title": "CHAPTER 1", "sub": "The Soil Frontier",
     "text": "Shallow soil beneath a meadow.\nOxygen seeps from above. Nitrate flows\nthrough the root zone. Methane bubbles up\nfrom decomposing organic matter.\n\nA good place to learn the ways\nof microbial life.\n\nArrows/WASD: Move\nSHIFT: Ride flow (planktonic)\nSPACE: Divide\nQ: Science Mode"},
    {"title": "CHAPTER 2", "sub": "The Deep Sediment",
     "text": "Ocean floor sediment. Hundreds of meters\nbelow the waves. Light never reaches here.\nOxygen is gone. Only iron and manganese\nremain as electron acceptors.\n\nFood is scarce. Every molecule counts.\nDiffusion dominates transport."},
    {"title": "CHAPTER 3", "sub": "The Methane Seeps",
     "text": "Active methane vents push CH4 through\nfractured rock. Sulfate mingles with\nthe rising gas. Toxic zones where\nreactive species damage cells.\n\nThis is where methanotrophs shine.\nConsume the methane. Save the climate."},
    {"title": "CHAPTER 4", "sub": "The Permafrost Edge",
     "text": "Warming temperatures thaw ancient ice.\nTrapped methane erupts in massive pulses.\nThe flow is fast and chaotic.\n\nThis is urgent. Every molecule of CH4\nthat escapes warms the planet further.\nCatch it before it rises."},
    {"title": "CHAPTER 5", "sub": "The Hydrothermal Realm",
     "text": "Deep ocean vents blast superheated fluid\nthrough mineral chimneys. Extreme chemistry.\nSulfate, iron, manganese - the full\nredox ladder in one place.\n\nOnly the strongest colonies survive here.\nThis is your final stand."},
]

# Level titles
level_titles = [
    "First Breath", "Roots of Life",
    "Into the Depths", "The Hungry Dark",
    "The Methane Vents", "Vent Guardians",
    "Thawing Grounds", "The Great Thaw",
    "The Abyss", "Earth's Last Stand",
]

# Level goals
level_goals = [3, 5, 4, 6, 5, 8, 6, 8, 8, 12]

# Chapter mapping: which levels belong to which chapter
chapter_levels = [
    (0, 1),   # Chapter 1: levels 0, 1
    (2, 3),   # Chapter 2: levels 2, 3
    (4, 5),   # Chapter 3: levels 4, 5
    (6, 7),   # Chapter 4: levels 6, 7
    (8, 9),   # Chapter 5: levels 8, 9
]

chapter_section_nums = [4, 5, 6, 7, 8]

# ═══════════════════════════════════════════════════════════════════════════
# SECTIONS 4-8: CHAPTERS 1-5
# ═══════════════════════════════════════════════════════════════════════════
for ch_idx in range(5):
    lev_start, lev_end = chapter_levels[ch_idx]
    sec_num = chapter_section_nums[ch_idx]
    intro = world_intros[ch_idx]

    doc.add_page_break()
    doc.add_heading(f'{sec_num}. {intro["title"]}: {intro["sub"]}', level=1)

    # Chapter Introduction
    doc.add_heading(f'{sec_num}.1  {intro["title"]} Introduction', level=2)
    doc.add_paragraph(
        f'This introduction screen is displayed when the player enters {intro["sub"]} '
        f'for the first time.'
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(intro["title"])
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(intro["sub"])
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    for tline in intro["text"].split("\n"):
        p = doc.add_paragraph(tline)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

    doc.add_paragraph()

    # Levels in this chapter
    sub_section = 2
    for lev_idx in range(lev_start, lev_end + 1):
        lev_num = lev_idx + 1
        lev_title = level_titles[lev_idx]

        # Mission Briefing
        doc.add_heading(f'{sec_num}.{sub_section}  Level {lev_num} – {lev_title} (Mission Briefing)', level=2)
        doc.add_paragraph(
            f'Shown before Level {lev_num} begins. The Elder briefs ARKE on the upcoming mission.'
        )
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run('MISSION BRIEFING')
        run.bold = True
        run.font.color.rgb = RGBColor(79, 163, 240)
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run(f'Level {lev_num}: {lev_title}')
        run.bold = True
        run.font.color.rgb = RGBColor(200, 170, 50)
        doc.add_paragraph()

        for line in mission_briefs[lev_idx]:
            speaker = line["speaker"]
            text = line["text"]

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)

            if speaker == "ELDER":
                run = p.add_run('ARCHAEON PRIME:')
                run.bold = True
                run.font.color.rgb = RGBColor(70, 70, 210)
            elif speaker == "ARKE":
                run = p.add_run('ARKE:')
                run.bold = True
                run.font.color.rgb = RGBColor(40, 207, 176)

            for tline in text.split("\n"):
                p2 = doc.add_paragraph(tline)
                p2.paragraph_format.left_indent = Inches(0.5)
                p2.paragraph_format.space_after = Pt(1)
                p2.paragraph_format.space_before = Pt(0)

        doc.add_paragraph()
        doc.add_paragraph(f'Goal: Place {level_goals[lev_idx]} colonies')
        doc.add_paragraph()
        sub_section += 1

        # Level Complete
        doc.add_heading(f'{sec_num}.{sub_section}  Level {lev_num} – Level Complete', level=2)
        doc.add_paragraph('Displayed when the player places enough colonies to meet the level goal.')
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('LEVEL COMPLETE!')
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(94, 207, 94)
        doc.add_paragraph()
        doc.add_paragraph('Substrates eaten: [count]')
        doc.add_paragraph('Methane consumed: [count]')
        doc.add_paragraph('N2O prevented: [count]')
        doc.add_paragraph('Colonies placed: [count]')
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('Score: [points]')
        run.bold = True
        doc.add_paragraph()
        doc.add_paragraph('PRESS ENTER')
        doc.add_paragraph()
        sub_section += 1

        # Science Discovery
        fact = science_facts[lev_idx]
        doc.add_heading(f'{sec_num}.{sub_section}  Level {lev_num} – Science Discovery: {fact["title"]}', level=2)
        doc.add_paragraph('Shown after the Level Complete screen. Teaches real science relevant to the level.')
        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run('SCIENCE DISCOVERY')
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(79, 163, 240)

        p = doc.add_paragraph()
        run = p.add_run(fact["title"])
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(200, 170, 50)

        doc.add_paragraph()
        for tline in fact["text"].split("\n"):
            p = doc.add_paragraph(tline)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)

        doc.add_paragraph()
        doc.add_paragraph('PRESS ENTER')
        doc.add_paragraph()
        sub_section += 1


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 9: VICTORY SCREEN
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('9. Victory Screen', level=1)
doc.add_paragraph(
    'Displayed after completing all 10 levels. Text appears progressively with timed reveals.'
)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('VICTORY!')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(200, 170, 50)
doc.add_paragraph()

doc.add_paragraph('Arke has colonized the underground!')
doc.add_paragraph()
doc.add_paragraph('From a single archaea to a thriving biofilm,')
doc.add_paragraph('you prevented greenhouse gases from escaping.')
doc.add_paragraph()
doc.add_paragraph('Methane consumed: [total count] molecules')
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Final Score: [total points]')
run.bold = True
run.font.size = Pt(16)
doc.add_paragraph()

doc.add_paragraph('Based on CompLaB3D')
doc.add_paragraph('Pore-Scale Biogeochemical Reactive Transport')
doc.add_paragraph()
doc.add_paragraph('University of Georgia')
doc.add_paragraph()
doc.add_paragraph('PRESS ENTER')

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 10: GAME OVER SCREEN
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('10. Game Over Screen', level=1)
doc.add_paragraph(
    'Displayed when the player\'s health reaches zero. One of the following messages '
    'is randomly shown.'
)
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('MICROBE LOST')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(240, 69, 69)
doc.add_paragraph()

doc.add_paragraph('Random death messages (one is selected each time):')
death_msgs = [
    "Starvation: the ultimate substrate limitation.",
    "The geochemical environment was too hostile.",
    "Without nutrients, even archaea perish.",
    "Position matters - poor pore connectivity is fatal.",
]
for msg in death_msgs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(f'- "{msg}"')
    run.italic = True

doc.add_paragraph()
doc.add_paragraph('Score: [accumulated points]')
doc.add_paragraph()
doc.add_paragraph('PRESS ENTER TO RETRY')

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 11: TUTORIAL HINTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('11. Tutorial Hints', level=1)
doc.add_paragraph(
    'These contextual hints appear at the top of the screen during the first two levels, '
    'triggered by player actions and game state.'
)
doc.add_paragraph()

tutorial_hints = [
    {"condition": "No substrates eaten yet (after 2 seconds)",
     "text": "Move with ARROWS/WASD - touch substrates to eat them!"},
    {"condition": "1-2 substrates eaten, health below 80",
     "text": "Eating restores HP and Energy - keep moving to find food!"},
    {"condition": "Health below 40",
     "text": "Health is low! Find substrates quickly or you'll starve!"},
    {"condition": "Growth bar full, no colonies placed yet",
     "text": "Growth bar full! Press SPACE to place a colony!"},
    {"condition": "After placing first colony, goal not yet reached",
     "text": "Colony placed! Keep eating to grow and place [goal] total."},
    {"condition": "Can ride flow, not currently riding, eaten 5+ substrates",
     "text": "Hold SHIFT to ride the flow current (planktonic mode)!"},
]

for hint in tutorial_hints:
    p = doc.add_paragraph()
    run = p.add_run(f'Condition: ')
    run.bold = True
    run = p.add_run(hint["condition"])
    run.italic = True

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(f'"{hint["text"]}"')
    doc.add_paragraph()

# Additional in-game text
doc.add_heading('In-Game Warning Text', level=2)
doc.add_paragraph()
doc.add_paragraph('When health drops below 30:')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"STARVING!"')
run.bold = True
run.font.color.rgb = RGBColor(240, 50, 50)
doc.add_paragraph()

doc.add_paragraph('When growth bar is full and ready to divide:')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"SPACE!"')
run.bold = True
run.font.color.rgb = RGBColor(200, 200, 50)
doc.add_paragraph()

doc.add_paragraph('When riding flow:')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"[RIDING FLOW]"')
run.bold = True
run.font.color.rgb = RGBColor(79, 163, 240)
doc.add_paragraph()

doc.add_paragraph('When Science Mode is active:')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"[SCIENCE MODE: Q]"')
run.bold = True
run.font.color.rgb = RGBColor(79, 163, 240)
doc.add_paragraph()

doc.add_paragraph('Flow riding hint (when available, not dividing):')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
run = p.add_run('"SHIFT: ride flow"')

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 12: PAUSE MENU
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('12. Pause Menu', level=1)
doc.add_paragraph('Accessed by pressing ESC during gameplay.')
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('PAUSED')
run.bold = True
run.font.size = Pt(18)
doc.add_paragraph()

doc.add_paragraph('> Resume')
doc.add_paragraph('  Mute / Unmute')
doc.add_paragraph('  Quit')
doc.add_paragraph()
doc.add_paragraph('(Arrow keys to navigate, ENTER to select)')

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 13: HUD ELEMENTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('13. HUD Elements and In-Game Labels', level=1)
doc.add_paragraph('Permanently visible during gameplay at the bottom of the screen.')
doc.add_paragraph()

doc.add_heading('Resource Bars', level=2)
doc.add_paragraph('HP  [===========] (Health - green to red based on value)')
doc.add_paragraph('EN  [===========] (Energy - blue)')
doc.add_paragraph('GR  [===========] (Growth - green, flashes yellow when full)')
doc.add_paragraph()

doc.add_heading('Status Text', level=2)
doc.add_paragraph('Colony: [current]/[goal]')
doc.add_paragraph('Score: [points]')
doc.add_paragraph('CH4: [methane eaten count]')
doc.add_paragraph('Lv[level number]')
doc.add_paragraph()

doc.add_heading('Flow Direction Indicator', level=2)
doc.add_paragraph('IN  >>>  OUT')
doc.add_paragraph('(Animated arrows showing water flow direction from inlet to outlet)')
doc.add_paragraph()
doc.add_paragraph('If level has toxic zones:')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('[purple square] = TOXIC (H2S)')
doc.add_paragraph()

doc.add_heading('Redox Ladder Display', level=2)
doc.add_paragraph('Shows available substrates for the current level with energy values:')
doc.add_paragraph('(Example for Chapter 1)')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('[blue]   O2    +20')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('[green]  NO3-  +15')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run('[red]    CH4   +5')
doc.add_paragraph()

doc.add_heading('Floating Popup Text', level=2)
doc.add_paragraph('When eating a substrate, a floating "+[energy] [formula]" text appears above the player and drifts upward.')
doc.add_paragraph('Examples:')
doc.add_paragraph('  +20 O2')
doc.add_paragraph('  +15 NO3-')
doc.add_paragraph('  +5 CH4')

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 14: SUBSTRATE INFORMATION
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('14. Substrate Information (Redox Ladder)', level=1)
doc.add_paragraph(
    'The six substrate types in the game, based on real thermodynamic free energy '
    'of microbial metabolisms. Listed in order of descending energy yield.'
)
doc.add_paragraph()

substrates = [
    ("Oxygen", "O2", "Blue", 20, 15, -818.0, "Aerobic respiration"),
    ("Nitrate", "NO3-", "Green", 15, 18, -649.0, "Denitrification - prevents N2O"),
    ("Manganese", "Mn(IV)", "Purple", 12, 12, -558.0, "Manganese reduction"),
    ("Iron", "Fe(III)", "Orange", 10, 10, -334.0, "Iron reduction"),
    ("Sulfate", "SO42-", "Yellow", 6, 8, -152.0, "Sulfate reduction"),
    ("Methane", "CH4", "Red", 5, 25, -31.0, "Methanotrophy - prevents warming!"),
]

table = doc.add_table(rows=1, cols=7)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Name'
hdr[1].text = 'Formula'
hdr[2].text = 'Color'
hdr[3].text = 'Energy'
hdr[4].text = 'Growth'
hdr[5].text = 'dG (kJ/mol)'
hdr[6].text = 'Description'
for h in hdr:
    for p in h.paragraphs:
        for r in p.runs:
            r.bold = True

for s in substrates:
    row = table.add_row().cells
    row[0].text = s[0]
    row[1].text = s[1]
    row[2].text = s[2]
    row[3].text = str(s[3])
    row[4].text = str(s[4])
    row[5].text = str(s[5])
    row[6].text = s[6]

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 15: SCIENCE PARAMETERS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('15. Science Parameters (from CompLaB3D)', level=1)
doc.add_paragraph('Real biogeochemical parameters used in the game, based on the CompLaB3D simulation framework.')
doc.add_paragraph()

params = [
    ("mu_max", "1.0", "Max specific growth rate [1/s]"),
    ("Ks", "1.0e-5", "Half-saturation constant [mol/L]"),
    ("Y", "0.4", "Yield coefficient [-]"),
    ("k_decay", "1.0e-9", "Decay rate [1/s]"),
    ("B_max", "100.0", "Max biomass density [kg/m3]"),
    ("D_pore", "1.0e-9", "Diffusion coefficient in pore [m2/s]"),
    ("D_biofilm", "2.0e-10", "Diffusion in biofilm [m2/s]"),
    ("visc_ratio", "10.0", "Viscosity ratio in biofilm"),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Parameter'
hdr[1].text = 'Value'
hdr[2].text = 'Description'
for h in hdr:
    for p in h.paragraphs:
        for r in p.runs:
            r.bold = True

for param in params:
    row = table.add_row().cells
    row[0].text = param[0]
    row[1].text = param[1]
    row[2].text = param[2]

doc.add_paragraph()
doc.add_heading('Game Balance Parameters', level=2)

balance = [
    ("move_speed", "2.5", "Tiles per second (normal movement)"),
    ("flow_ride_mult", "2.0", "Speed multiplier when riding flow"),
    ("max_health", "100.0", "Maximum health points"),
    ("max_energy", "100.0", "Maximum energy"),
    ("max_growth", "100.0", "Maximum growth (division threshold)"),
    ("health_decay", "1.5", "HP lost per second (starvation)"),
    ("energy_decay", "0.8", "Energy lost per second"),
    ("toxic_damage", "20.0", "Damage per second in toxic zone"),
    ("division_cost", "100.0", "Growth needed to divide"),
    ("substrate_spawn", "1.2", "Seconds between substrate spawns"),
    ("score_eat", "10", "Points per substrate eaten"),
    ("score_colony", "100", "Points per colony placed"),
    ("score_level", "1000", "Points per level completed"),
    ("score_methane", "25", "Bonus points for eating CH4"),
    ("score_no3", "15", "Bonus points for eating NO3 (prevents N2O)"),
]

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Parameter'
hdr[1].text = 'Value'
hdr[2].text = 'Description'
for h in hdr:
    for p in h.paragraphs:
        for r in p.runs:
            r.bold = True

for param in balance:
    row = table.add_row().cells
    row[0].text = param[0]
    row[1].text = param[1]
    row[2].text = param[2]

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 16: ENVIRONMENT DESCRIPTIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('16. Environment Descriptions', level=1)
doc.add_paragraph('The five unique environments in the game, each based on real geological settings.')
doc.add_paragraph()

envs = [
    ("The Soil Frontier", "Where Life Begins",
     "Shallow soil environment. Open pores with good oxygen availability. "
     "The tutorial area where players learn basic mechanics."),
    ("The Deep Sediment", "Darkness Below",
     "Ocean floor sediment. No oxygen. Maze-like passages with iron and manganese "
     "as primary electron acceptors. Dark and scarce environment."),
    ("The Methane Seeps", "Rivers of Fire",
     "Active methane vents with fractured rock. Sulfate and methane are primary substrates. "
     "Toxic zones (H2S) appear for the first time."),
    ("The Permafrost Edge", "The Melting World",
     "Thawing arctic permafrost. Fast, chaotic flow. Ancient trapped methane erupts. "
     "Urgent gameplay with high flow speeds."),
    ("The Hydrothermal Realm", "Earth's Furnace",
     "Deep-sea hydrothermal vents. Extreme chemistry with the full redox ladder available. "
     "Maximum difficulty with toxic zones, fast flow, and many rivals."),
]

for name, sub, desc in envs:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}')
    run.bold = True
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    run = p.add_run(f'"{sub}"')
    run.italic = True
    doc.add_paragraph(desc)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 17: LEVEL DEFINITIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('17. Level Definitions', level=1)
doc.add_paragraph('Complete configuration for all 10 levels.')
doc.add_paragraph()

levels = [
    {"num": 1, "title": "First Breath", "ch": "The Soil Frontier", "env": 0,
     "map": "30x20", "goal": 3, "porosity": "70%", "subs": "O2, NO3, CH4",
     "flow": 0.6, "toxic": "0%", "rivals": 0},
    {"num": 2, "title": "Roots of Life", "ch": "The Soil Frontier", "env": 0,
     "map": "35x22", "goal": 5, "porosity": "65%", "subs": "O2, NO3, CH4",
     "flow": 0.7, "toxic": "0%", "rivals": 1},
    {"num": 3, "title": "Into the Depths", "ch": "The Deep Sediment", "env": 1,
     "map": "35x25", "goal": 4, "porosity": "50%", "subs": "NO3, Fe(III), CH4",
     "flow": 0.3, "toxic": "0%", "rivals": 2},
    {"num": 4, "title": "The Hungry Dark", "ch": "The Deep Sediment", "env": 1,
     "map": "40x28", "goal": 6, "porosity": "45%", "subs": "NO3, Mn(IV), Fe(III), CH4",
     "flow": 0.25, "toxic": "0%", "rivals": 3},
    {"num": 5, "title": "The Methane Vents", "ch": "The Methane Seeps", "env": 2,
     "map": "35x22", "goal": 5, "porosity": "60%", "subs": "SO4, CH4",
     "flow": 0.5, "toxic": "15%", "rivals": 2},
    {"num": 6, "title": "Vent Guardians", "ch": "The Methane Seeps", "env": 2,
     "map": "40x25", "goal": 8, "porosity": "55%", "subs": "SO4, CH4, Fe(III)",
     "flow": 0.6, "toxic": "20%", "rivals": 3},
    {"num": 7, "title": "Thawing Grounds", "ch": "The Permafrost Edge", "env": 3,
     "map": "40x25", "goal": 6, "porosity": "55%", "subs": "O2, NO3, CH4",
     "flow": 0.8, "toxic": "10%", "rivals": 3},
    {"num": 8, "title": "The Great Thaw", "ch": "The Permafrost Edge", "env": 3,
     "map": "45x28", "goal": 8, "porosity": "50%", "subs": "NO3, SO4, CH4",
     "flow": 1.0, "toxic": "15%", "rivals": 4},
    {"num": 9, "title": "The Abyss", "ch": "The Hydrothermal Realm", "env": 4,
     "map": "45x25", "goal": 8, "porosity": "55%", "subs": "SO4, Fe(III), Mn(IV), CH4",
     "flow": 1.2, "toxic": "20%", "rivals": 4},
    {"num": 10, "title": "Earth's Last Stand", "ch": "The Hydrothermal Realm", "env": 4,
     "map": "50x30", "goal": 12, "porosity": "50%", "subs": "SO4, Fe(III), Mn(IV), CH4, NO3",
     "flow": 1.5, "toxic": "25%", "rivals": 5},
]

table = doc.add_table(rows=1, cols=9)
table.style = 'Table Grid'
hdr = table.rows[0].cells
headers = ['Level', 'Title', 'Chapter', 'Map', 'Goal', 'Porosity', 'Substrates', 'Toxic', 'Rivals']
for i, h in enumerate(headers):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

for lev in levels:
    row = table.add_row().cells
    row[0].text = str(lev["num"])
    row[1].text = lev["title"]
    row[2].text = lev["ch"]
    row[3].text = lev["map"]
    row[4].text = str(lev["goal"])
    row[5].text = lev["porosity"]
    row[6].text = lev["subs"]
    row[7].text = lev["toxic"]
    row[8].text = str(lev["rivals"])
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════
output_path = '/home/user/Complab3D-Biotic-Kinetic-AndersopnNR/ARKE_Game_Dialogues_and_Information.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
